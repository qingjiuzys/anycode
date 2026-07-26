#!/usr/bin/env python3
"""Generate brand template.dotx for any brand kit."""
from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(2)

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO / "brand-kits" / "lib"))
from brand_kit import color_from_tokens, find_brand_kit, load_tokens  # noqa: E402


def ensure_heading_styles(doc: Document, tokens: dict) -> None:
    primary = RGBColor(*color_from_tokens(tokens, "primary"))
    styles = doc.styles
    for level, size in [(1, 22), (2, 16), (3, 13)]:
        name = f"Heading {level}"
        try:
            st = styles[name]
        except KeyError:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = tokens.get("fonts", {}).get("heading_en", "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = primary


def apply_header_footer(doc: Document, tokens: dict) -> None:
    primary = RGBColor(*color_from_tokens(tokens, "primary"))
    muted = RGBColor(*color_from_tokens(tokens, "muted_text"))
    for section in doc.sections:
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        hp.text = tokens.get("name", "Document")
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = primary
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer = (tokens.get("footer") or {}).get("default_text") or f"Confidential · {tokens.get('name', 'Brand')}"
        fp.text = footer
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = muted


def main() -> int:
    brand = sys.argv[1] if len(sys.argv) > 1 else "fde-editorial"
    kit = find_brand_kit(brand)
    tokens = load_tokens(kit)
    out_dir = kit / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "template.dotx"

    doc = Document()
    ensure_heading_styles(doc, tokens)
    apply_header_footer(doc, tokens)
    doc.add_heading("Document Title", level=1)
    doc.add_heading("Section", level=2)
    doc.add_paragraph("Body text placeholder.")
    tmp = out_dir / "template.docx"
    doc.save(str(tmp))
    tmp.replace(out)
    meta = {"brand_kit": brand, "path": str(out), "styles": ["Heading 1", "Heading 2", "Heading 3"]}
    (out_dir / "template.meta.json").write_text(json.dumps(meta, indent=2) + "\n", encoding="utf-8")
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
