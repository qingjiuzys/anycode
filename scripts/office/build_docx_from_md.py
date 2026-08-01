#!/usr/bin/env python3
"""Build a branded .docx from Markdown (lingqi brand-kit: headings, header/footer)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO / "brand-kits" / "lib"))
from brand_kit import color_from_tokens, emit_artifact, find_brand_kit, load_tokens  # noqa: E402

HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
BULLET_RE = re.compile(r"^[-*]\s+(.+)$")


def apply_header_footer(doc: Document, tokens: dict, title: str):
    primary = RGBColor(*color_from_tokens(tokens, "primary"))
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = title[:80]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = primary
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = tokens.get("name", "Document") + " · Confidential"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*color_from_tokens(tokens, "muted_text"))


def style_heading(paragraph, level: int, tokens: dict):
    primary = RGBColor(*color_from_tokens(tokens, "primary"))
    for run in paragraph.runs:
        run.font.color.rgb = primary
        run.font.bold = True
        run.font.size = Pt({1: 22, 2: 16, 3: 13}.get(level, 13))


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: run report.md [output.docx] [brand_kit]", file=sys.stderr)
        return 1
    md_path = Path(sys.argv[1]).resolve()
    out_path = (
        Path(sys.argv[2]).resolve()
        if len(sys.argv) > 2
        else md_path.with_suffix(".docx")
    )
    brand = sys.argv[3] if len(sys.argv) > 3 else "fde-editorial"
    tokens = load_tokens(find_brand_kit(brand))
    kit = find_brand_kit(brand)
    dotx = kit / "docx" / "template.dotx"
    if not dotx.is_file():
        import subprocess

        subprocess.run([sys.executable, str(REPO / "scripts" / "office" / "build_brand_dotx.py"), brand], check=False)
    text = md_path.read_text(encoding="utf-8")
    if not text.strip():
        print("ERROR: empty markdown", file=sys.stderr)
        return 1

    doc = Document()
    doc_title = md_path.stem.replace("-", " ").replace("_", " ")
    apply_header_footer(doc, tokens, doc_title)

    saw_heading = False
    saw_decision_action = False
    h1_count = 0
    h2_count = 0
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        hm = HEADING_RE.match(line)
        if hm:
            level = len(hm.group(1))
            title = hm.group(2).strip()
            p = doc.add_heading(title, level=min(level, 3))
            style_heading(p, level, tokens)
            saw_heading = True
            if level == 1:
                h1_count += 1
            elif level == 2:
                h2_count += 1
            continue
        bm = BULLET_RE.match(line)
        body = bm.group(1).strip() if bm else line.strip()
        lower = body.lower()
        if lower.startswith("decision:") or lower.startswith("action:"):
            saw_decision_action = True
        p = doc.add_paragraph(body)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = tokens["fonts"].get("body_en", "Calibri")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), tokens["fonts"].get("body_zh", "Songti SC"))

    if not saw_heading or h1_count < 1:
        print("ERROR: need at least one H1 heading", file=sys.stderr)
        return 1
    if h2_count < 2 and len(doc.paragraphs) > 20:
        print("WARN: book-style doc should have multiple H2 sections", file=sys.stderr)
    if not saw_decision_action:
        print(
            "ERROR: need at least one Decision: or Action: line with owner/date",
            file=sys.stderr,
        )
        return 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    emit_artifact(out_path, "document")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
